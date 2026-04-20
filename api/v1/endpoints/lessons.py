import io
import json
import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

import docx
import pypdf
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from Backend.api.deps import get_current_user
from Backend.db.session import get_db
from Backend.models.user import User
from Backend.models.user_lesson import UserLesson
from Backend.schemas.lesson import (
    LessonGenerationResponse,
    LessonNavigation,
    LessonPage,
    LessonProgressSnapshot,
    LessonProgressUpdateRequest,
    LessonSaveRequest,
    LessonSourceDocument,
    SavedLessonDetail,
    SavedLessonSummary,
)
from Backend.services.workflows import ExerciseOrchestrator
from Backend.services.workflows.config import (
    ContentGenerationRequest,
    ContentLevel,
    ContentType,
    StudentProfile,
)
from Backend.services.web_research_agent import get_web_research_agent
from Backend.services.workflows.utils.grounding import build_source_context, chunk_source_material, normalize_text

logger = logging.getLogger(__name__)
router = APIRouter()

SUPPORTED_DOCUMENT_TYPES = {".pdf", ".docx", ".txt", ".md"}
ALLOWED_LEARNING_STYLES = {"visual", "auditory", "kinesthetic", "reading/writing"}
ALLOWED_LEARNING_PACES = {"slow", "normal", "fast"}
ALLOWED_QUIZ_TYPES = {"multiple_choice", "fill_blank", "true_false"}


def _parse_list_form_field(raw_value: str | None) -> list[str]:
    if raw_value is None:
        return []

    cleaned = raw_value.strip()
    if not cleaned:
        return []

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, list):
            return [str(item).strip() for item in parsed if str(item).strip()]
    except json.JSONDecodeError:
        pass

    return [part.strip() for part in re.split(r"[\n,;]+", cleaned) if part.strip()]


def _derive_topic_from_prompt(prompt: str) -> str:
    normalized = " ".join(prompt.split())
    if not normalized:
        return "Custom Lesson"

    for separator in [".", "?", "!"]:
        end_index = normalized.find(separator)
        if 5 <= end_index <= 80:
            return normalized[:end_index].strip()

    return normalized[:80].rstrip(" .,:;-") or "Custom Lesson"


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    chunks: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            chunks.append(page_text)
    return "\n".join(chunks)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def _extract_document_text(file_name: str, file_bytes: bytes) -> tuple[str, str]:
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_TYPES:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Supported types: {', '.join(sorted(SUPPORTED_DOCUMENT_TYPES))}"
        )

    if suffix == ".pdf":
        text = _extract_text_from_pdf(file_bytes)
    elif suffix == ".docx":
        text = _extract_text_from_docx(file_bytes)
    else:
        text = file_bytes.decode("utf-8")

    cleaned = text.strip()

    if suffix == ".pdf" and not cleaned:
        raise ValueError(
            "No selectable text could be extracted from this PDF. Text-based PDFs are supported, but scanned/image PDFs and embedded images/formulas are not currently OCR'd."
        )

    return cleaned, suffix


def _to_json_safe(value: Any) -> Any:
    if hasattr(value, "value"):
        return value.value

    if isinstance(value, dict):
        return {str(key): _to_json_safe(item) for key, item in value.items()}

    if isinstance(value, list):
        return [_to_json_safe(item) for item in value]

    return value


def _extract_agent_payload(workflow_log: list[dict[str, Any]], agent_name: str) -> dict[str, Any] | None:
    for entry in workflow_log:
        if entry.get("agent") == agent_name and entry.get("success"):
            payload = entry.get("data")
            if isinstance(payload, dict):
                return payload
    return None


def _question_signature(text: str) -> str:
    return re.sub(r"\s+", " ", normalize_text(text)).casefold()


def _pick_progression_stage(index: int, total: int) -> str:
    if total <= 1:
        return "foundation"
    if index == 0:
        return "foundation"
    if index >= total - 1:
        return "integration"
    if index == total - 2:
        return "advanced"
    return "core"


def _first_sentence(value: str) -> str:
    normalized = normalize_text(value)
    if not normalized:
        return ""

    segments = re.split(r"(?<=[.!?])\s+", normalized)
    return (segments[0] if segments else normalized).strip()


def _extract_candidate_term(statement: str, fallback: str) -> str:
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9_-]{3,}", statement or "")
    for token in tokens:
        lowered = token.casefold()
        if lowered in {
            "this",
            "that",
            "which",
            "where",
            "when",
            "from",
            "with",
            "into",
            "about",
            "while",
            "their",
            "there",
            "because",
            "through",
        }:
            continue
        if fallback and lowered in fallback.casefold():
            continue
        return token

    fallback_parts = re.findall(r"[A-Za-z][A-Za-z0-9_-]+", fallback or "")
    return fallback_parts[0] if fallback_parts else "the core concept"


def _build_did_you_know_fact(section: dict[str, Any], topic: str) -> str:
    source_support = section.get("source_support") or []
    key_points = section.get("key_points") or []
    section_title = normalize_text(str(section.get("title", "")))

    candidate = ""
    if source_support:
        candidate = _first_sentence(str(source_support[0]))
    elif key_points:
        candidate = _first_sentence(str(key_points[0]))
    elif section_title:
        candidate = f"{section_title} is a key part of understanding {topic}."

    if not candidate:
        candidate = f"Understanding this section strengthens your grasp of {topic}."

    if len(candidate) < 35:
        candidate = f"{candidate} This is often overlooked but highly useful in practice."

    return candidate


def _build_quick_check(section: dict[str, Any], topic: str, order_index: int) -> dict[str, Any]:
    section_title = normalize_text(str(section.get("title", ""))) or "this section"
    source_support = section.get("source_support") or []
    key_points = section.get("key_points") or []

    evidence = ""
    if source_support:
        evidence = _first_sentence(str(source_support[0]))
    elif key_points:
        evidence = _first_sentence(str(key_points[0]))
    if not evidence:
        evidence = f"{section_title} explains a practical principle in {topic}."

    if order_index % 2 == 0:
        distractors = [
            f"{section_title} is unrelated to {topic} and has no practical value.",
            f"{section_title} should always be ignored when solving {topic} problems.",
        ]
        options = [evidence, *distractors]
        return {
            "type": "quick_check",
            "purpose": "immediate_understanding",
            "format": "multiple_choice",
            "question": f"Quick check: Which statement best reflects {section_title}?",
            "options": options,
            "correct_answer_index": 0,
            "explanation": "Choose the option that matches the lesson explanation in this section.",
        }

    answer_term = _extract_candidate_term(evidence, section_title)
    blanked = re.sub(re.escape(answer_term), "_____", evidence, count=1, flags=re.IGNORECASE)
    if blanked == evidence:
        blanked = f"{section_title} helps explain _____ in {topic}."

    return {
        "type": "quick_check",
        "purpose": "immediate_understanding",
        "format": "fill_blank",
        "question": f"Quick check: Fill in the blank. {blanked}",
        "correct_answers": [answer_term],
        "explanation": "Fill the blank with the core term emphasized in this section.",
    }


def _enrich_sections_with_in_lesson_interactions(
    *,
    topic: str,
    sections: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], set[str]]:
    enriched_sections: list[dict[str, Any]] = []
    in_lesson_elements: list[dict[str, Any]] = []
    blocked_signatures: set[str] = set()

    main_sections = [section for section in sections if str(section.get("type")) == "main_content"]
    total_main_sections = max(1, len(main_sections))
    main_index = 0

    for section in sections:
        section_copy = dict(section)

        if str(section_copy.get("type")) != "main_content":
            enriched_sections.append(section_copy)
            continue

        stage = _pick_progression_stage(main_index, total_main_sections)
        did_you_know = {
            "type": "did_you_know",
            "purpose": "engagement",
            "prompt": "Did you know?",
            "fact": _build_did_you_know_fact(section_copy, topic),
        }
        quick_check = _build_quick_check(section_copy, topic, main_index)

        section_copy["progression_stage"] = stage
        section_copy["interactive_elements"] = [did_you_know, quick_check]

        in_lesson_elements.append(
            {
                "section_id": section_copy.get("id"),
                "section_title": section_copy.get("title"),
                "progression_stage": stage,
                "did_you_know": did_you_know,
                "quick_check": quick_check,
                "purpose": "in_lesson_engagement_and_comprehension",
            }
        )

        blocked_signatures.add(_question_signature(str(quick_check.get("question", ""))))
        main_index += 1
        enriched_sections.append(section_copy)

    return enriched_sections, in_lesson_elements, blocked_signatures


def _to_mcq_from_true_false(question_item: dict[str, Any]) -> dict[str, Any]:
    statement = normalize_text(str(question_item.get("question", "")))
    correct_bool = bool(question_item.get("correct_answer", False))
    correct_option = "True" if correct_bool else "False"
    return {
        "format": "multiple_choice",
        "question": f"True or False: {statement}",
        "options": [correct_option, "False" if correct_bool else "True"],
        "correct_answer_index": 0,
        "explanation": normalize_text(str(question_item.get("explanation", "")))
        or "Evaluate whether the statement matches the lesson context.",
    }


def _build_application_exercise(
    *,
    topic: str,
    section: dict[str, Any],
    order_index: int,
) -> dict[str, Any]:
    title = normalize_text(str(section.get("title", ""))) or "the concept"
    key_points = section.get("key_points") or []
    anchor = _first_sentence(str(key_points[0])) if key_points else f"the core principle of {title}"

    if order_index % 2 == 0:
        options = [
            f"Prioritize {anchor} while considering the context and constraints.",
            f"Ignore {title} and focus only on memorizing definitions.",
            f"Apply {title} the same way in every scenario without evaluation.",
        ]
        return {
            "format": "multiple_choice",
            "question": f"Application practice: In a real {topic} scenario, what is the best first move for {title}?",
            "options": options,
            "correct_answer_index": 0,
            "explanation": "Application tasks reward context-aware decisions, not rote steps.",
        }

    answer_term = _extract_candidate_term(anchor, title)
    question = f"Application practice: When using {title}, start by evaluating _____ before acting."
    return {
        "format": "fill_blank",
        "question": question,
        "correct_answers": [answer_term],
        "explanation": "Applying the concept begins with selecting the right contextual signal.",
    }


def _build_end_of_lesson_exercises(
    *,
    topic: str,
    sections: list[dict[str, Any]],
    flashcard_payload: dict[str, Any] | None,
    quiz_payload: dict[str, Any] | None,
    blocked_signatures: set[str],
) -> dict[str, Any]:
    flashcards = list((flashcard_payload or {}).get("flashcards", []) or [])
    quiz_questions = list((quiz_payload or {}).get("questions", []) or [])

    used_signatures = set(blocked_signatures)

    def is_available(question_text: str) -> bool:
        signature = _question_signature(question_text)
        if not signature or signature in used_signatures:
            return False
        used_signatures.add(signature)
        return True

    memorization: list[dict[str, Any]] = []
    consolidation: list[dict[str, Any]] = []
    application: list[dict[str, Any]] = []

    for card in flashcards:
        card_question = normalize_text(str(card.get("question", "")))
        card_answer = normalize_text(str(card.get("answer", "")))
        if not card_question or not card_answer:
            continue
        if not is_available(card_question):
            continue

        options = [
            card_answer,
            f"A loosely related detail about {topic}.",
            f"A statement that contradicts the main lesson idea.",
        ]
        memorization.append(
            {
                "exercise_id": f"mem-{len(memorization) + 1}",
                "goal": "memorization",
                "purpose": "retention",
                "difficulty": "easy",
                "format": "multiple_choice",
                "question": card_question,
                "options": options,
                "correct_answer_index": 0,
                "explanation": "Memorization exercises reinforce core facts and definitions.",
            }
        )

        if len(memorization) >= 6:
            break

    for question_item in quiz_questions:
        qtype = str(question_item.get("type", "")).casefold()
        question_text = normalize_text(str(question_item.get("question", "")))
        if not question_text:
            continue
        if not is_available(question_text):
            continue

        if qtype == "multiple_choice":
            consolidation.append(
                {
                    "exercise_id": f"con-{len(consolidation) + 1}",
                    "goal": "consolidation",
                    "purpose": "understanding",
                    "difficulty": "intermediate",
                    "format": "multiple_choice",
                    "question": question_text,
                    "options": list(question_item.get("options", []) or []),
                    "correct_answer_index": int(question_item.get("correct_answer", 0)),
                    "explanation": normalize_text(str(question_item.get("explanation", ""))),
                }
            )
        elif qtype == "fill_blank":
            consolidation.append(
                {
                    "exercise_id": f"con-{len(consolidation) + 1}",
                    "goal": "consolidation",
                    "purpose": "understanding",
                    "difficulty": "intermediate",
                    "format": "fill_blank",
                    "question": question_text,
                    "correct_answers": list(question_item.get("correct_answers", []) or []),
                    "explanation": normalize_text(str(question_item.get("explanation", ""))),
                }
            )
        elif qtype == "true_false":
            converted = _to_mcq_from_true_false(question_item)
            consolidation.append(
                {
                    "exercise_id": f"con-{len(consolidation) + 1}",
                    "goal": "consolidation",
                    "purpose": "understanding",
                    "difficulty": "intermediate",
                    **converted,
                }
            )

        if len(consolidation) >= 8:
            break

    main_sections = [section for section in sections if str(section.get("type")) == "main_content"]
    for index, section in enumerate(main_sections[:6]):
        synthesized = _build_application_exercise(topic=topic, section=section, order_index=index)
        if not is_available(str(synthesized.get("question", ""))):
            continue

        application.append(
            {
                "exercise_id": f"app-{len(application) + 1}",
                "goal": "application",
                "purpose": "practice",
                "difficulty": "advanced" if index >= 2 else "intermediate",
                **synthesized,
            }
        )

        if len(application) >= 6:
            break

    exercise_sets = [
        {
            "set_id": "memorization",
            "title": "Memorization",
            "purpose": "retention",
            "recommended_difficulty": "easy",
            "exercises": memorization,
        },
        {
            "set_id": "consolidation",
            "title": "Consolidation",
            "purpose": "deeper understanding",
            "recommended_difficulty": "intermediate",
            "exercises": consolidation,
        },
        {
            "set_id": "application",
            "title": "Application and Practice",
            "purpose": "skill transfer and problem solving",
            "recommended_difficulty": "intermediate_to_advanced",
            "exercises": application,
        },
    ]

    flattened = [
        exercise
        for exercise_set in exercise_sets
        for exercise in exercise_set["exercises"]
    ]

    return {
        "exercise_sets": exercise_sets,
        "all_exercises": flattened,
        "total_exercises": len(flattened),
        "progression": ["easy", "intermediate", "advanced"],
        "design_rules": {
            "in_lesson_prompts": "immediate understanding and engagement",
            "end_of_lesson_exercises": "retention, consolidation, and application",
            "overlap_avoidance": "Question signatures are deduplicated between stages",
        },
    }


def _build_lesson_pages(
    topic: str,
    prompt: str,
    source_documents: list[LessonSourceDocument],
    lesson_payload: dict[str, Any] | None,
    flashcard_payload: dict[str, Any] | None,
    quiz_payload: dict[str, Any] | None,
    coding_payload: dict[str, Any] | None,
    mindmap_payload: dict[str, Any] | None,
    include_answer_key: bool,
    external_sources: list[dict[str, Any]] | None = None,
) -> list[LessonPage]:
    pages: list[LessonPage] = []
    order = 1
    external_sources = external_sources or []

    pages.append(
        LessonPage(
            page_id="overview",
            order=order,
            page_type="overview",
            title="Lesson Overview",
            description="Prompt and source context used to generate this lesson.",
            estimated_time_minutes=2,
            data={
                "prompt": prompt,
                "source_documents": [_to_json_safe(doc.model_dump()) for doc in source_documents],
                "external_sources": _to_json_safe(external_sources),
            },
        )
    )
    order += 1

    enriched_sections: list[dict[str, Any]] = []
    in_lesson_elements: list[dict[str, Any]] = []
    in_lesson_blocked_signatures: set[str] = set()

    if lesson_payload:
        lesson_data = lesson_payload.get("lesson", {})
        sections = list(lesson_payload.get("sections", []) or [])
        (
            enriched_sections,
            in_lesson_elements,
            in_lesson_blocked_signatures,
        ) = _enrich_sections_with_in_lesson_interactions(
            topic=topic,
            sections=sections,
        )
        resources = lesson_payload.get("learning_resources", [])
        pages.append(
            LessonPage(
                page_id="theory",
                order=order,
                page_type="theory",
                title="Theory and Knowledge",
                description="Progressive textbook-style explanations with in-lesson interaction prompts.",
                estimated_time_minutes=int(lesson_payload.get("estimated_duration_minutes", 0)),
                data={
                    "lesson": _to_json_safe(lesson_data),
                    "sections": _to_json_safe(enriched_sections),
                    "total_sections": int(lesson_payload.get("total_sections", len(enriched_sections))),
                    "in_lesson_interactions": _to_json_safe(in_lesson_elements),
                    "module_design": {
                        "style": "textbook_progressive",
                        "theory_progression": [
                            section.get("progression_stage")
                            for section in enriched_sections
                            if isinstance(section, dict) and section.get("progression_stage")
                        ],
                        "interaction_policy": "did_you_know_and_quick_checks_interleaved",
                    },
                },
            )
        )
        order += 1

        if resources or external_sources:
            pages.append(
                LessonPage(
                    page_id="resources",
                    order=order,
                    page_type="resources",
                    title="Learning Resources",
                    description="Supplementary references for deeper learning.",
                    estimated_time_minutes=5,
                    data={
                        "items": _to_json_safe(resources),
                        "external_sources": _to_json_safe(external_sources),
                    },
                )
            )
            order += 1

    if external_sources and not any(page.page_id == "resources" for page in pages):
        pages.append(
            LessonPage(
                page_id="resources",
                order=order,
                page_type="resources",
                title="Learning Resources",
                description="Supplementary references for deeper learning.",
                estimated_time_minutes=5,
                data={
                    "items": [],
                    "external_sources": _to_json_safe(external_sources),
                },
            )
        )
        order += 1

    exercise_plan = _build_end_of_lesson_exercises(
        topic=topic,
        sections=enriched_sections,
        flashcard_payload=flashcard_payload,
        quiz_payload=quiz_payload,
        blocked_signatures=in_lesson_blocked_signatures,
    )

    if flashcard_payload:
        flashcards = [
            card
            for card in (flashcard_payload.get("flashcards", []) or [])
            if _question_signature(str(card.get("question", ""))) not in in_lesson_blocked_signatures
        ]
        pages.append(
            LessonPage(
                page_id="flashcards",
                order=order,
                page_type="flashcards",
                title="End-of-Lesson Memorization Deck",
                description="Post-lesson memorization cards for retention practice.",
                estimated_time_minutes=10,
                data={
                    "cards": _to_json_safe(flashcards),
                    "total_cards": int(flashcard_payload.get("total_cards", len(flashcards))),
                    "purpose": "post_lesson_memorization",
                },
            )
        )
        order += 1

    if quiz_payload or exercise_plan.get("total_exercises", 0) > 0:
        questions = [_to_json_safe(question) for question in exercise_plan.get("all_exercises", [])]

        pages.append(
            LessonPage(
                page_id="quiz",
                order=order,
                page_type="quiz",
                title="End-of-Lesson Exercises",
                description="Structured practice for memorization, consolidation, and application.",
                estimated_time_minutes=max(
                    int((quiz_payload or {}).get("estimated_duration_minutes", 0)),
                    max(10, len(questions) * 2),
                ),
                data={
                    "quiz": _to_json_safe((quiz_payload or {}).get("quiz", {})),
                    "questions": questions,
                    "total_questions": len(questions),
                    "question_types": _to_json_safe((quiz_payload or {}).get("question_types", [])),
                    "exercise_sets": _to_json_safe(exercise_plan.get("exercise_sets", [])),
                    "exercise_progression": _to_json_safe(exercise_plan.get("progression", [])),
                    "design_rules": _to_json_safe(exercise_plan.get("design_rules", {})),
                    "answer_key_included": include_answer_key,
                },
            )
        )
        order += 1

    if coding_payload:
        coding_tasks = [_to_json_safe(task) for task in coding_payload.get("coding_tasks", [])]
        include_coding = bool(coding_payload.get("include_coding_exercises", False))

        if include_coding and coding_tasks:
            pages.append(
                LessonPage(
                    page_id="coding",
                    order=order,
                    page_type="coding",
                    title="Coding Exercises",
                    description="Hands-on coding practice with server-evaluated test cases.",
                    estimated_time_minutes=max(10, len(coding_tasks) * 12),
                    data={
                        "tasks": coding_tasks,
                        "total_tasks": int(coding_payload.get("total_tasks", len(coding_tasks))),
                        "language": coding_payload.get("language"),
                        "language_id": coding_payload.get("language_id"),
                        "decision_reason": coding_payload.get("decision_reason"),
                    },
                )
            )
            order += 1

    if mindmap_payload:
        json_format = mindmap_payload.get("json_format")
        parsed_json: Any = json_format
        if isinstance(json_format, str):
            try:
                parsed_json = json.loads(json_format)
            except json.JSONDecodeError:
                parsed_json = json_format

        pages.append(
            LessonPage(
                page_id="mindmap",
                order=order,
                page_type="mindmap",
                title="Mind Map",
                description="Visual map of concept relationships.",
                estimated_time_minutes=8,
                data={
                    "mindmap": _to_json_safe(mindmap_payload.get("mindmap_structure", {})),
                    "visualization": _to_json_safe(parsed_json),
                },
            )
        )

    return pages


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []

    for value in values:
        if value and value not in seen:
            seen.add(value)
            result.append(value)

    return result


def _validate_navigation_consistency(
    pages: list[LessonPage],
    navigation: LessonNavigation,
) -> set[str]:
    page_ids = [page.page_id for page in pages]
    unique_page_ids = set(page_ids)

    if len(unique_page_ids) != len(page_ids):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Each lesson page must have a unique page_id",
        )

    if navigation.total_pages != len(pages):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="navigation.total_pages must match the number of pages",
        )

    if navigation.total_pages != len(navigation.page_order):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="navigation.page_order length must match navigation.total_pages",
        )

    if set(navigation.page_order) != unique_page_ids:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="navigation.page_order must include exactly the same page_ids as pages",
        )

    if navigation.default_page_id not in unique_page_ids:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="navigation.default_page_id must reference a valid page_id",
        )

    return unique_page_ids


def _normalize_completed_page_ids(
    completed_page_ids: list[str],
    valid_page_ids: set[str],
) -> list[str]:
    normalized = _dedupe_preserve_order(completed_page_ids)
    invalid_page_ids = [page_id for page_id in normalized if page_id not in valid_page_ids]

    if invalid_page_ids:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "completed_page_ids contains invalid page ids: "
                + ", ".join(invalid_page_ids)
            ),
        )

    return normalized


def _validate_current_page_id(current_page_id: str | None, valid_page_ids: set[str]) -> str | None:
    if current_page_id is None:
        return None

    if current_page_id not in valid_page_ids:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="current_page_id must reference a valid page_id",
        )

    return current_page_id


def _build_progress_snapshot(
    current_page_id: str | None,
    completed_page_ids: list[str],
    total_pages: int,
) -> LessonProgressSnapshot:
    completed_count = min(total_pages, len(set(completed_page_ids)))
    progress_percent = round((completed_count / total_pages) * 100, 2) if total_pages > 0 else 0.0

    return LessonProgressSnapshot(
        current_page_id=current_page_id,
        completed_page_ids=completed_page_ids,
        total_pages=total_pages,
        progress_percent=progress_percent,
    )


def _serialize_saved_lesson(
    lesson: UserLesson,
    include_content: bool,
) -> SavedLessonSummary | SavedLessonDetail:
    pages = [LessonPage.model_validate(page) for page in (lesson.pages or [])]
    navigation = LessonNavigation.model_validate(lesson.navigation or {})
    valid_page_ids = set(navigation.page_order)

    completed_page_ids = [
        page_id
        for page_id in _dedupe_preserve_order(lesson.completed_page_ids or [])
        if page_id in valid_page_ids
    ]

    current_page_id = lesson.current_page_id if lesson.current_page_id in valid_page_ids else navigation.default_page_id
    progress = _build_progress_snapshot(
        current_page_id=current_page_id,
        completed_page_ids=completed_page_ids,
        total_pages=navigation.total_pages,
    )

    base_payload = {
        "id": lesson.id,
        "title": lesson.title,
        "topic": lesson.topic,
        "created_at": lesson.created_at,
        "updated_at": lesson.updated_at,
        "last_opened_at": lesson.last_opened_at,
        "progress": progress,
    }

    if not include_content:
        return SavedLessonSummary(**base_payload)

    return SavedLessonDetail(
        **base_payload,
        prompt=lesson.prompt,
        pages=pages,
        navigation=navigation,
        source_documents=[
            LessonSourceDocument.model_validate(item)
            for item in (lesson.source_documents or [])
        ],
        execution_summary=lesson.execution_summary or {},
        quality_metrics=lesson.quality_metrics or {},
        workflow_issues=lesson.workflow_issues or [],
    )


async def _get_user_lesson_or_404(
    db: AsyncSession,
    lesson_id: UUID,
    user_id: UUID,
) -> UserLesson:
    result = await db.execute(
        select(UserLesson).where(
            UserLesson.id == lesson_id,
            UserLesson.user_id == user_id,
        )
    )
    lesson = result.scalar_one_or_none()

    if lesson is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Lesson not found")

    return lesson


@router.post("/generate", response_model=LessonGenerationResponse)
async def generate_lesson(
    prompt: str = Form(..., min_length=10),
    topic: str | None = Form(default=None),
    subject: str | None = Form(default=None),
    subtopics: str | None = Form(default=None),
    learning_objectives: str | None = Form(default=None),
    current_level: str = Form(default="intermediate"),
    learning_style: str = Form(default="reading/writing"),
    learning_pace: str = Form(default="normal"),
    daily_study_time_minutes: int = Form(default=30, ge=5, le=300),
    max_quiz_questions: int = Form(default=10, ge=1, le=50),
    quiz_question_types: str | None = Form(default=None),
    include_mindmap: bool = Form(default=False),
    include_coding_exercises: bool | None = Form(default=None),
    include_answer_key: bool = Form(default=False),
    include_external_sources: bool = Form(default=False),
    external_search_query: str | None = Form(default=None),
    max_external_sources: int = Form(default=6, ge=1, le=12),
    files: list[UploadFile] | None = File(default=None),
    current_user: User = Depends(get_current_user),
):
    """
    Generate a multi-page lesson from user prompt and optional uploaded documents.

    The response contains a pages array for frontend navigation between lesson sections
    (overview, theory, flashcards, quiz, and optional mindmap/resources).
    """
    files = files or []

    if learning_style not in ALLOWED_LEARNING_STYLES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"learning_style must be one of: {', '.join(sorted(ALLOWED_LEARNING_STYLES))}",
        )

    if learning_pace not in ALLOWED_LEARNING_PACES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"learning_pace must be one of: {', '.join(sorted(ALLOWED_LEARNING_PACES))}",
        )

    try:
        level = ContentLevel(current_level)
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="current_level must be one of: beginner, intermediate, advanced",
        ) from exc

    requested_quiz_types = _parse_list_form_field(quiz_question_types)
    if requested_quiz_types:
        invalid_quiz_types = [item for item in requested_quiz_types if item not in ALLOWED_QUIZ_TYPES]
        if invalid_quiz_types:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Invalid quiz_question_types: {', '.join(invalid_quiz_types)}",
            )
    else:
        requested_quiz_types = ["multiple_choice", "fill_blank", "true_false"]

    source_documents: list[LessonSourceDocument] = []
    source_materials: list[str] = []
    external_sources: list[dict[str, str]] = []
    for uploaded_file in files:
        if not uploaded_file.filename:
            continue

        file_bytes = await uploaded_file.read()
        if not file_bytes:
            continue

        try:
            extracted_text, extension = _extract_document_text(uploaded_file.filename, file_bytes)
        except UnicodeDecodeError as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to decode {uploaded_file.filename}. Ensure the file is UTF-8 text.",
            ) from exc
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to process file {uploaded_file.filename}: {str(exc)}",
            ) from exc

        source_documents.append(
            LessonSourceDocument(
                file_name=uploaded_file.filename,
                file_type=extension,
                extracted_characters=len(extracted_text),
                excerpt=extracted_text[:600] if extracted_text else None,
            )
        )
        if extracted_text:
            source_materials.extend(
                chunk_source_material(
                    extracted_text,
                    chunk_size=2600,
                    overlap=260,
                    max_chunks=260,
                )
            )

    resolved_topic = topic.strip() if topic and topic.strip() else _derive_topic_from_prompt(prompt)
    resolved_subtopics = _parse_list_form_field(subtopics)
    if not resolved_subtopics:
        resolved_subtopics = [
            resolved_topic,
            "Core Concepts",
            "Practical Applications",
        ]

    resolved_objectives = _parse_list_form_field(learning_objectives)
    if not resolved_objectives:
        resolved_objectives = [
            f"Understand the theory of {resolved_topic}",
            f"Memorize key terms and definitions in {resolved_topic}",
            f"Assess comprehension through a guided quiz",
        ]

    if include_external_sources:
        search_seed = external_search_query.strip() if external_search_query and external_search_query.strip() else prompt
        research_agent = get_web_research_agent()
        external_sources = await research_agent.research(
            topic=resolved_topic,
            prompt=search_seed,
            subtopics=resolved_subtopics,
            learning_objectives=resolved_objectives,
            max_sources=max_external_sources,
        )
        source_materials.extend(research_agent.to_source_materials(external_sources))

    if source_documents:
        resolved_objectives.append("Incorporate context from uploaded documents in explanations and examples")

    if external_sources:
        resolved_objectives.append("Use relevant external references to improve accuracy and completeness")

    lesson_content_types = [ContentType.LESSON, ContentType.FLASHCARD, ContentType.QUIZ]
    coding_decision_mode = "auto"
    if include_coding_exercises is not False:
        lesson_content_types.append(ContentType.CODING_TASK)
        coding_decision_mode = "force" if include_coding_exercises is True else "auto"
    else:
        coding_decision_mode = "skip"
    if include_mindmap:
        lesson_content_types.append(ContentType.MINDMAP)

    student_profile = StudentProfile(
        student_id=str(current_user.id),
        name=current_user.full_name or current_user.email,
        subject=subject or resolved_topic,
        current_level=level,
        learning_style=learning_style,
        knowledge_gaps=[],
        strengths=[],
        learning_pace=learning_pace,
        preferred_content_types=lesson_content_types,
        daily_study_time_minutes=daily_study_time_minutes,
    )

    workflow_request = ContentGenerationRequest(
        student_profile=student_profile,
        topic=resolved_topic,
        subtopics=resolved_subtopics,
        learning_objectives=resolved_objectives,
        content_types=lesson_content_types,
        prompt=prompt,
        source_materials=source_materials,
        source_context=build_source_context(prompt, source_materials),
        max_items=max_quiz_questions,
        quiz_question_types=requested_quiz_types,
        coding_decision_mode=coding_decision_mode,
    )

    try:
        orchestrator = ExerciseOrchestrator()
        orchestrator_result = await orchestrator.run({"request": workflow_request})
    except Exception as exc:
        logger.exception("Failed to run lesson generation workflow")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(exc),
        ) from exc

    if not orchestrator_result.get("success"):
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=orchestrator_result.get("error", "Lesson workflow failed"),
        )

    workflow_data = orchestrator_result.get("data", {})
    workflow_log = workflow_data.get("workflow_log", [])

    lesson_payload = _extract_agent_payload(workflow_log, "LessonCreatorAgent")
    flashcard_payload = _extract_agent_payload(workflow_log, "FlashcardCreatorAgent")
    quiz_payload = _extract_agent_payload(workflow_log, "QuizCreatorAgent")
    coding_payload = _extract_agent_payload(workflow_log, "CodingTaskCreatorAgent")
    mindmap_payload = _extract_agent_payload(workflow_log, "MindmapCreatorAgent")

    pages = _build_lesson_pages(
        topic=resolved_topic,
        prompt=prompt,
        source_documents=source_documents,
        lesson_payload=lesson_payload,
        flashcard_payload=flashcard_payload,
        quiz_payload=quiz_payload,
        coding_payload=coding_payload,
        mindmap_payload=mindmap_payload,
        include_answer_key=include_answer_key,
        external_sources=external_sources,
    )

    if len(pages) <= 1:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Lesson generation did not produce content pages",
        )

    page_order = [page.page_id for page in pages]
    navigation = LessonNavigation(
        total_pages=len(pages),
        page_order=page_order,
        default_page_id=page_order[0],
    )

    quality_metrics = _to_json_safe(workflow_data.get("quality_metrics", {}))

    return LessonGenerationResponse(
        success=True,
        lesson_id=str(uuid4()),
        title=f"{resolved_topic} Learning Lesson",
        topic=resolved_topic,
        prompt=prompt,
        pages=pages,
        navigation=navigation,
        source_documents=source_documents,
        execution_summary=_to_json_safe(workflow_data.get("execution_summary", {})),
        quality_metrics=quality_metrics,
        workflow_issues=quality_metrics.get("quality_issues", []) if isinstance(quality_metrics, dict) else [],
    )


@router.post("/save", response_model=SavedLessonDetail, status_code=status.HTTP_201_CREATED)
async def save_lesson(
    payload: LessonSaveRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Persist a generated lesson so the user can reopen it later."""
    valid_page_ids = _validate_navigation_consistency(payload.pages, payload.navigation)

    completed_page_ids = _normalize_completed_page_ids(
        payload.completed_page_ids,
        valid_page_ids,
    )

    current_page_id = payload.current_page_id or payload.navigation.default_page_id
    current_page_id = _validate_current_page_id(current_page_id, valid_page_ids)

    now = datetime.now(timezone.utc)
    user_lesson = UserLesson(
        user_id=current_user.id,
        title=payload.title,
        topic=payload.topic,
        prompt=payload.prompt,
        pages=[page.model_dump() for page in payload.pages],
        navigation=payload.navigation.model_dump(),
        source_documents=[doc.model_dump() for doc in payload.source_documents],
        execution_summary=payload.execution_summary,
        quality_metrics=payload.quality_metrics,
        workflow_issues=payload.workflow_issues,
        current_page_id=current_page_id,
        completed_page_ids=completed_page_ids,
        last_opened_at=now,
    )

    db.add(user_lesson)
    await db.commit()
    await db.refresh(user_lesson)

    return _serialize_saved_lesson(user_lesson, include_content=True)


@router.get("", response_model=list[SavedLessonSummary])
async def list_saved_lessons(
    limit: int = 20,
    offset: int = 0,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List persisted lessons for the current user, newest first."""
    safe_limit = max(1, min(limit, 100))
    safe_offset = max(0, offset)

    result = await db.execute(
        select(UserLesson)
        .where(UserLesson.user_id == current_user.id)
        .order_by(UserLesson.updated_at.desc())
        .offset(safe_offset)
        .limit(safe_limit)
    )

    lessons = result.scalars().all()
    return [
        _serialize_saved_lesson(lesson, include_content=False)
        for lesson in lessons
    ]


@router.get("/{lesson_id:uuid}", response_model=SavedLessonDetail)
async def get_saved_lesson(
    lesson_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get a persisted lesson in full detail for rendering and resume."""
    lesson = await _get_user_lesson_or_404(db, lesson_id, current_user.id)
    lesson.last_opened_at = datetime.now(timezone.utc)

    await db.commit()
    await db.refresh(lesson)

    return _serialize_saved_lesson(lesson, include_content=True)


@router.patch("/{lesson_id:uuid}/progress", response_model=SavedLessonDetail)
async def update_saved_lesson_progress(
    lesson_id: UUID,
    payload: LessonProgressUpdateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update resume progress (current page and completed pages) for a saved lesson."""
    lesson = await _get_user_lesson_or_404(db, lesson_id, current_user.id)
    navigation = LessonNavigation.model_validate(lesson.navigation)
    valid_page_ids = set(navigation.page_order)

    if "completed_page_ids" in payload.model_fields_set and payload.completed_page_ids is not None:
        lesson.completed_page_ids = _normalize_completed_page_ids(payload.completed_page_ids, valid_page_ids)

    if "current_page_id" in payload.model_fields_set:
        if payload.current_page_id is None:
            lesson.current_page_id = navigation.default_page_id
        else:
            lesson.current_page_id = _validate_current_page_id(payload.current_page_id, valid_page_ids)

    if lesson.current_page_id is None:
        lesson.current_page_id = navigation.default_page_id

    lesson.last_opened_at = datetime.now(timezone.utc)

    await db.commit()
    await db.refresh(lesson)

    return _serialize_saved_lesson(lesson, include_content=True)


@router.get("/health")
async def lessons_health() -> dict[str, Any]:
    """Health check for lesson-generation endpoint and orchestrator wiring."""
    return {
        "status": "healthy",
        "service": "lessons-api",
        "orchestrator": "ExerciseOrchestrator",
    }
